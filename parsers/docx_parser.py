from docx import Document

def parse_docx(path: str) -> str:
    """Extract text from DOCX file including tables"""
    try:
        doc = Document(path)
        parts = []
        
        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        
        return "\n".join(parts)
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")
