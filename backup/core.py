# =============================================================================
# COMPLETE RESUME EXTRACTION SYSTEM
# =============================================================================

# -----------------------------------------------------------------------------
# schemas/resume.py - Data Models
# -----------------------------------------------------------------------------
from pydantic import BaseModel
from typing import List, Optional

class ExperienceItem(BaseModel):
    company: Optional[str] = None
    role: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    description: Optional[str] = None

class EducationItem(BaseModel):
    institution: Optional[str] = None
    degree: Optional[str] = None
    start_year: Optional[str] = None
    end_year: Optional[str] = None

class Resume(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    links: List[str] = []
    summary: Optional[str] = None
    skills: List[str] = []
    experience: List[ExperienceItem] = []
    education: List[EducationItem] = []
    certifications: List[str] = []


# -----------------------------------------------------------------------------
# parsers/pdf_parser.py - PDF Text Extraction
# -----------------------------------------------------------------------------
import pdfplumber

def parse_pdf(path: str) -> str:
    """Extract text from PDF file"""
    text = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                text.append(t)
        return "\n".join(text)
    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")


# -----------------------------------------------------------------------------
# parsers/docx_parser.py - DOCX Text Extraction
# -----------------------------------------------------------------------------
from docx import Document

def parse_docx(path: str) -> str:
    """Extract text from DOCX file including tables"""
    try:
        doc = Document(path)
        parts = []
        
        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        
        return "\n".join(parts)
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")


# -----------------------------------------------------------------------------
# parsers/image_parser.py - OCR for Images
# -----------------------------------------------------------------------------
import pytesseract
from PIL import Image, ImageOps, ImageFilter

def parse_image(path: str) -> str:
    """Extract text from image using OCR"""
    try:
        img = Image.open(path)
        # Preprocess for better OCR
        img = ImageOps.grayscale(img)
        img = img.filter(ImageFilter.SHARPEN)
        text = pytesseract.image_to_string(img)
        return text
    except Exception as e:
        raise Exception(f"Error parsing image: {str(e)}")


# -----------------------------------------------------------------------------
# extractors/patterns.py - Regex-based Extraction
# -----------------------------------------------------------------------------
import re

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[-\s]?)?(?:\d{10}|\d{3}[-\s]?\d{3}[-\s]?\d{4})")
URL_RE = re.compile(r"https?://[^\s]+|(?:www\.)?linkedin\.com/[^\s]+|github\.com/[^\s]+")

def extract_email(text: str) -> Optional[str]:
    """Extract first email address"""
    m = EMAIL_RE.search(text)
    return m.group(0) if m else None

def extract_phone(text: str) -> Optional[str]:
    """Extract first phone number"""
    m = PHONE_RE.search(text)
    return m.group(0) if m else None

def extract_links(text: str) -> List[str]:
    """Extract all URLs"""
    return list(set(URL_RE.findall(text)))


# -----------------------------------------------------------------------------
# extractors/nlp.py - NER with spaCy
# -----------------------------------------------------------------------------
import spacy

try:
    nlp = spacy.load("en_core_web_sm")
except:
    print("Warning: spaCy model not found. Run: python -m spacy download en_core_web_sm")
    nlp = None

def extract_entities(text: str) -> dict:
    """Extract named entities using spaCy"""
    if not nlp:
        return {"PERSON": [], "ORG": [], "GPE": [], "DATE": []}
    
    doc = nlp(text[:100000])  # Limit text size for performance
    ents = {"PERSON": [], "ORG": [], "GPE": [], "DATE": []}
    
    for e in doc.ents:
        if e.label_ in ents:
            ents[e.label_].append(e.text)
    
    return ents


# -----------------------------------------------------------------------------
# extractors/sections.py - Section Splitting
# -----------------------------------------------------------------------------
HEADERS = [
    "summary", "objective", "profile", "about",
    "education", "academic", "qualification",
    "experience", "work experience", "employment", "work history",
    "skills", "technical skills", "core competencies", "expertise",
    "projects", "certifications", "certificates", "licenses"
]

def split_sections(text: str) -> dict:
    """Split resume into sections based on headers"""
    lines = text.splitlines()
    sections = {}
    current = "other"
    
    for line in lines:
        line_lower = line.strip().lower()
        
        # Check if line is a header
        matched = False
        for header in HEADERS:
            if line_lower.startswith(header) or line_lower == header:
                current = header
                sections[current] = []
                matched = True
                break
        
        if not matched and line.strip():
            sections.setdefault(current, []).append(line)
    
    return {k: "\n".join(v).strip() for k, v in sections.items() if v}


# -----------------------------------------------------------------------------
# extractors/skills.py - Skills Extraction
# -----------------------------------------------------------------------------
KNOWN_SKILLS = {
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "ruby", "php", "swift", "kotlin", "go", "rust",
    "react", "angular", "vue", "node.js", "express", "django", "flask", "spring", "asp.net",
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "git", "ci/cd",
    "machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch",
    "html", "css", "sass", "bootstrap", "tailwind",
    "rest api", "graphql", "microservices", "agile", "scrum"
}

def extract_skills(section_text: str) -> List[str]:
    """Extract skills from skills section"""
    if not section_text:
        return []
    
    tokens = re.split(r"[,\n;•·\|]", section_text.lower())
    found_skills = set()
    
    for token in tokens:
        cleaned = token.strip()
        if cleaned in KNOWN_SKILLS:
            found_skills.add(cleaned)
    
    return sorted(list(found_skills))


# -----------------------------------------------------------------------------
# extractors/education.py - Education Extraction
# -----------------------------------------------------------------------------
DEGREE_WORDS = [
    "b.tech", "b.e", "b.e.", "btech", "bachelor",
    "bsc", "b.sc", "ba", "b.a",
    "m.tech", "m.e", "m.e.", "mtech", "master",
    "msc", "m.sc", "ms", "m.s", "mba", "m.b.a",
    "phd", "ph.d", "doctorate",
    "diploma", "associate"
]

def extract_education(section_text: str) -> List[dict]:
    """Extract education entries"""
    if not section_text:
        return []
    
    items = []
    blocks = section_text.split("\n\n")
    
    for block in blocks:
        lines = [l for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        
        block_lower = block.lower()
        has_degree = any(deg in block_lower for deg in DEGREE_WORDS)
        
        if has_degree:
            years = re.findall(r"(19|20)\d{2}", block)
            degree = next((d for d in DEGREE_WORDS if d in block_lower), None)
            
            items.append({
                "institution": lines[0] if lines else None,
                "degree": degree,
                "start_year": years[0] if len(years) >= 1 else None,
                "end_year": years[1] if len(years) >= 2 else years[0] if len(years) == 1 else None
            })
    
    return items


# -----------------------------------------------------------------------------
# extractors/experience.py - Experience Extraction
# -----------------------------------------------------------------------------
DATE_RANGE = re.compile(
    r"([A-Za-z]{3,9}\s?\d{4}|\d{4})\s?[-–—]\s?([A-Za-z]{3,9}\s?\d{4}|\d{4}|present|current)",
    re.IGNORECASE
)

def extract_experience(section_text: str) -> List[dict]:
    """Extract work experience entries"""
    if not section_text:
        return []
    
    items = []
    blocks = section_text.split("\n\n")
    
    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        
        # Find date range
        m = DATE_RANGE.search(block)
        dates = (m.group(1), m.group(2)) if m else (None, None)
        
        # Parse first line for role and company
        first = lines[0]
        role, company = None, None
        
        if " at " in first:
            parts = first.split(" at ", 1)
            role = parts[0].strip()
            company = parts[1].strip()
        elif " - " in first and len(lines) > 1:
            role = first
            company = lines[1] if len(lines) > 1 else None
        else:
            role = first
        
        # Description is remaining lines
        desc_lines = lines[1:] if company else lines[2:] if len(lines) > 2 else []
        description = "\n".join(desc_lines) if desc_lines else None
        
        items.append({
            "company": company,
            "role": role,
            "start_date": dates[0],
            "end_date": dates[1],
            "description": description
        })
    
    return items


# -----------------------------------------------------------------------------
# app/pipeline.py - Main Processing Pipeline
# -----------------------------------------------------------------------------
import os

def extract_text(path: str) -> str:
    """Extract text from file based on extension"""
    ext = os.path.splitext(path)[1].lower()
    
    if ext == ".pdf":
        return parse_pdf(path)
    elif ext in [".docx", ".doc"]:
        return parse_docx(path)
    elif ext in [".jpg", ".jpeg", ".png", ".tiff", ".bmp"]:
        return parse_image(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def to_json(text: str) -> dict:
    """Convert extracted text to structured JSON"""
    # Split into sections
    sections = split_sections(text)
    
    # Extract entities
    ents = extract_entities(text)
    
    # Extract skills
    skills_text = sections.get("skills", sections.get("technical skills", ""))
    skills = extract_skills(skills_text)
    
    # Extract education
    edu_text = sections.get("education", sections.get("academic", ""))
    education = extract_education(edu_text)
    
    # Extract experience
    exp_text = sections.get("experience", sections.get("work experience", ""))
    experience = extract_experience(exp_text)
    
    # Get summary
    summary = sections.get("summary", sections.get("objective", sections.get("profile", None)))
    
    # Extract basic info
    name = ents["PERSON"][0] if ents["PERSON"] else None
    location = ents["GPE"][0] if ents["GPE"] else None
    
    # Extract certifications
    cert_text = sections.get("certifications", sections.get("certificates", ""))
    certifications = [c.strip() for c in cert_text.splitlines() if c.strip()] if cert_text else []
    
    return {
        "name": name,
        "email": extract_email(text),
        "phone": extract_phone(text),
        "location": location,
        "links": extract_links(text),
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "education": education,
        "certifications": certifications
    }


# -----------------------------------------------------------------------------
# app/main.py - FastAPI Application
# -----------------------------------------------------------------------------
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
import tempfile
import os

app = FastAPI(
    title="Resume Extractor API",
    description="Extract structured data from resume files (PDF, DOCX, Images)",
    version="1.0.0"
)

@app.get("/")
async def root():
    return {
        "message": "Resume Extractor API",
        "endpoints": {
            "POST /extract": "Upload a resume file to extract data",
            "GET /health": "Check API health"
        }
    }

@app.get("/health")
async def health():
    return {"status": "healthy"}

@app.post("/extract")
async def extract_resume(file: UploadFile = File(...)):
    """
    Extract structured data from resume file
    
    Supports: PDF, DOCX, JPG, PNG
    """
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(400, "No file provided")
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            # Extract text
            text = extract_text(tmp_path)
            
            # Convert to structured JSON
            data = to_json(text)
            
            return {
                "status": "success",
                "filename": file.filename,
                "data": data
            }
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"Processing error: {str(e)}")


# -----------------------------------------------------------------------------
# Run the application
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)